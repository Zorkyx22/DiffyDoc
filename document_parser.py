import re
import docx

from datetime import datetime
from openai import OpenAI
from langchain_ollama import OllamaLLM
from tqdm import tqdm
from argparse import ArgumentParser

from constants import (END_FILE_PREFIX, SYNTAX_PROMPT as PROMPT, OPENAI_API_KEY, AI_NAME, DEFAULT_MODEL)
from insert_comment import add_comment_to_elements_in_place

from rich import print

from rich.progress import (
    BarColumn,
    Progress,
    TextColumn,
    TimeElapsedColumn,
    SpinnerColumn,
    track,
)

progress = Progress(
    TextColumn("[bold green]{task.fields[task_name]}", justify="right"),
    TextColumn("[bold red]{task.fields[error_message]}", justify="left"),
    BarColumn(bar_width=None),
    SpinnerColumn("dots"),
    TimeElapsedColumn(),
)


def handle_chunk(message_chunk):
    return agent.invoke(PROMPT.format(message_chunk))

def analyze_and_comment(filename):
    task_id = progress.add_task(f"Working on {filename}...", task_name="Initiating...", error_message="", start=True, total=None)
    doc = docx.Document(filename)
    total_paragraphs = len(doc.paragraphs)
    total_comments = 0
    for i, paragraph in enumerate(doc.paragraphs):
        progress.update(task_id, task_name=f"Handling paragraph {i}/{total_paragraphs}\nTotal number of commented paragraphs: {total_comments}")
        if(paragraph.text != "" and len(paragraph.runs) > 0 and not paragraph.style.name.startswith("Head")):
            comment = handle_chunk(paragraph.text)
            if not comment.startswith("This syntax is correct"):
                add_comment_to_elements_in_place(doc, [paragraph._element], AI_NAME, comment)
                total_comments+=1
    progress.update(task_id, task_name=f"Completed! Commented {total_comments} paragraphs out of {total_paragraphs} analyzed!", completed=True)
    print(f"Done!")
    doc.save("SyntAI_output.docx")

parser = ArgumentParser(description="A script utility that makes use of OpenAI GPT to work with docx documents without compromising its styling. You must have an Ollama server running and at least one model pulled. See README for instructions to setup.")
parser.add_argument("file", action="store", help="The docx file you wish to translate")
parser.add_argument("--model", "-m", action="store", default=DEFAULT_MODEL, help=f"The model to use with the local ollama server. The default is {DEFAULT_MODEL}")
if __name__=="__main__":
    args = parser.parse_args()
    agent = OllamaLLM(model=args.model)
    with progress:
        analyze_and_comment(args.file)
